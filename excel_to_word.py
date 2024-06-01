import pandas as pd
import random
from docx import Document
from docx.shared import RGBColor

df = pd.read_excel('sorular.xlsx')

# Word belgesi oluştur
doc = Document()

for index, row in df.iterrows():
    soru = row['SORU']
    dogru_cevap = row['DOGRU']
    yanlis_cevaplar = [row['YANLIŞ CEVAP 1'], row['YANLIŞ CEVAP 2'], row['YANLIŞ CEVAP 3'], row['YANLIŞ CEVAP 4']]
    
    cevaplar = yanlis_cevaplar + [dogru_cevap]
    random.shuffle(cevaplar)
    
    doc.add_paragraph(soru)
    for i, cevap in enumerate(cevaplar):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"{chr(65 + i)}) {cevap}")
        
        if cevap == dogru_cevap:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
    doc.add_paragraph("")  

doc.save('sorular.docx')
